from docx import Document

# Create a new Document
doc = Document()

# Add Title
doc.add_heading('MD. MONIRUZZAMAN', 0)

# Add Contact Information
doc.add_paragraph('North Kazipara, Mirpur-1216\n'
                  'C. 01645820113 | E. ammoniruzzaman@gmail.com\n')

# Add Professional Summary
doc.add_heading('Professional Summary', level=1)
doc.add_paragraph(
    "Dynamic Software Engineer with over 2 years of experience at Transcom Ltd. and 9 months at eAppyears, "
    "excelling in cross-platform application development for Android, iOS, desktop, and web using Flutter. "
    "Successfully led software development projects for renowned brands like KFC and Pizza Hut. Delivered impactful solutions "
    "across key initiatives, including a kiosk (self-order system), e-invoice, data lake, OKR, mReporting (For SKF, TDCL), "
    "Secondary sales (TCPL), and attendance systems. Skilled solution specialist and Python script writer with 6 months of "
    "backend development experience. Demonstrated expertise in driving innovation, optimizing workflows, and leading teams "
    "to achieve strategic goals."
)

# Add Experience
doc.add_heading('Experience', level=1)

# Software Engineer
doc.add_heading('Software Engineer', level=2)
doc.add_paragraph('Technology, Transcom Ltd. — January 2023 to Present')
doc.add_paragraph(
    "- Spearhead the development of cross-platform applications for Android, iOS, desktop, and web using Flutter, leveraging BLoC for efficient state management.\n"
    "- Currently leading the development of the KFC restaurant self-order kiosk system, enhancing customer experience and streamlining operations.\n"
    "- Focus on backend development for the last 6 months, enhancing server-side logic and database performance.\n"
    "- Engage in Python scripting and data lake integration throughout the year to optimize data workflows.\n"
    "- Lead software development projects for KFC and Pizza Hut, ensuring alignment with business objectives.\n"
    "- Designed and implemented major systems, including e-invoice platforms, OKR tracking, reporting, and attendance systems."
)

# Associate Software Engineer
doc.add_heading('Associate Software Engineer', level=2)
doc.add_paragraph('Technology, Transcom Ltd. — February 2023 to December 2023')
doc.add_paragraph(
    "- Dedicated to Flutter development, implementing features and UI components with BLoC state management for responsive applications.\n"
    "- Built and refined user interfaces, ensuring responsiveness and cross-platform compatibility.\n"
    "- Collaborated with teams to deliver feature-rich, user-friendly applications on time."
)

# Junior Software Engineer
doc.add_heading('Junior Software Engineer', level=2)
doc.add_paragraph('eAppyears — May 2022 to January 2023')
doc.add_paragraph(
    "- Focused on Android development using Flutter, contributing to a restaurant POS system for the UK.\n"
    "- Designed and implemented robust application features, ensuring seamless user experiences.\n"
    "- Debugged and resolved issues to improve application stability and performance.\n"
    "- Gained foundational expertise in cross-platform development and backend integration."
)

# Add Skills
doc.add_heading('Skills', level=1)
doc.add_paragraph(
    "- **Programming Languages:** Dart, Python, C, C++, JavaScript, SQL\n"
    "- **Frameworks & Technologies:** Flutter, Web2py, RESTful APIs, XML/JSON Parsing\n"
    "- **Development Expertise:** Cross-platform application development (Android, iOS, desktop, web)\n"
    "- **Software & Tools:** Android, Visual Studio Code, RealTerm, Git, Docker, Azure App Services\n"
    "- **Database Management:** MySQL, Oracle, PostgreSQL, Firebase\n"
    "- **Project Management:** Team leadership for KFC and Pizza Hut software development\n"
    "- **Key Domains:** Self-order kiosk systems, e-invoicing, data lake integration, OKRs, reporting, and attendance systems\n"
    "- **Problem-Solving:** Solution specialization and Python scripting for automation and backend logic\n"
    "- **Soft Skills:** Team collaboration, adaptability, attention to detail, and effective communication"
)

# Add Education
doc.add_heading('Education', level=1)
doc.add_paragraph('Bachelor of Engineering (Electronics and Communication Engineering)\n'
                  'National University — January 2016 to December 2020\n'
                  'CGPA: 3.64')

# Add Certifications
doc.add_heading('Certifications', level=1)
doc.add_paragraph('• Certified Mobile Application Development using Flutter')

# Add Languages
doc.add_heading('Languages', level=1)
doc.add_paragraph('• Bengali – Fluent\n• English – Proficient (Intermediate)')

# Add References
doc.add_heading('References', level=1)
doc.add_paragraph('Md. Moniruzzaman\nSoftware Engineer\nTechnology Division, Transcom Ltd\n'
                  'C. 01645820113 | E. ammoniruzzaman@gmail.com')

# Save to file
file_path = "/mnt/data/MD_Moniruzzaman_Resume.docx"
doc.save(file_path)

file_path
